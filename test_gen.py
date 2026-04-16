import docx
import pandas as pd
from pathlib import Path

out_dir = Path("test_docs")
out_dir.mkdir(exist_ok=True)

# Generate DOCX
doc = docx.Document()
doc.add_heading('Test Document', 0)
doc.add_paragraph('This is a simple paragraph with some ')
doc.paragraphs[-1].add_run('bold').bold = True
doc.paragraphs[-1].add_run(' and ')
doc.paragraphs[-1].add_run('italic').italic = True
doc.paragraphs[-1].add_run(' text.')

table = doc.add_table(rows=2, cols=2)
table.cell(0, 0).text = 'Header 1'
table.cell(0, 1).text = 'Header 2'
table.cell(1, 0).text = 'Data 1'
table.cell(1, 1).text = 'Data 2'

doc.save(out_dir / 'test.docx')

# Generate XLSX
df1 = pd.DataFrame({'A': [1, 2], 'B': [3, 4]})
df2 = pd.DataFrame({'Name': ['Alice', 'Bob'], 'Age': [25, 30]})

with pd.ExcelWriter(out_dir / 'test.xlsx') as writer:
    df1.to_excel(writer, sheet_name='Sheet1', index=False)
    df2.to_excel(writer, sheet_name='Sheet2', index=False)

# Not generating PDF as we need a heavier lib, but we can verify xlsx and docx first!
print("Generated test.docx and test.xlsx")
